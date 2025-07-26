import os
import pathlib
import traceback
import sys
from datetime import datetime
from flask import Flask, request, render_template, send_file, flash, redirect, url_for, jsonify
from werkzeug.utils import secure_filename
from docx import Document
from docx.shared import RGBColor
import zipfile
import tempfile

# Crear aplicación Flask
app = Flask(__name__)
app.secret_key = os.environ.get('SECRET_KEY', 'tu_clave_secreta_muy_segura_aqui')
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size

# Configuración de carpetas
UPLOAD_FOLDER = os.path.abspath('uploads')
PROCESSED_FOLDER = os.path.abspath('processed')
ALLOWED_EXTENSIONS = {'docx', 'docm'}

# Asegurar que las carpetas existen
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(PROCESSED_FOLDER, exist_ok=True)

# Configuración para producción
if os.environ.get('FLASK_ENV') == 'production':
    app.config.update(
        SECRET_KEY=os.environ.get('SECRET_KEY', 'clave-super-secreta-produccion'),
        DEBUG=False,
        TESTING=False
    )

def allowed_file(filename):
    """Verificar si el archivo tiene una extensión permitida"""
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def procesar_documento(doc_path, paso=70):
    """
    Procesa un documento Word insertando caracteres "0" invisibles
    cada 'paso' caracteres en párrafos desde la página 2.
    PRESERVA COMPLETAMENTE EL FORMATO ORIGINAL - Método Simplificado.
    
    Args:
        doc_path (str): Ruta al documento Word
        paso (int): Intervalo de caracteres para insertar los "0"
    
    Returns:
        str: Ruta del documento procesado
    """
    try:
        # Verificar archivo antes de abrir
        if not os.path.exists(doc_path):
            raise Exception(f"Archivo no encontrado: {doc_path}")
        
        file_size = os.path.getsize(doc_path)
        if file_size == 0:
            raise Exception("Archivo vacío")
        
        print(f"🔄 Procesando documento: {doc_path} ({file_size} bytes)")
        
        # Intentar abrir el documento con validación
        try:
            doc = Document(doc_path)
            print(f"✅ Documento abierto correctamente")
        except Exception as e:
            print(f"❌ Error al abrir documento: {e}")
            raise Exception(f"El archivo no es un documento Word válido: {str(e)}")
        
        # Verificar que el documento tiene contenido
        total_paragraphs = len(doc.paragraphs)
        print(f"📊 Total de párrafos: {total_paragraphs}")
        
        if total_paragraphs == 0:
            print("⚠️ Documento sin párrafos, procesando como vacío")
        
        # Procesar párrafos (saltando títulos y párrafos muy cortos)
        processed_count = 0
        for i, para in enumerate(doc.paragraphs):
            try:
                # Saltar párrafos vacíos o muy cortos (solo títulos de 1-2 palabras)
                if len(para.text.strip()) < 20:  # Cambio de 50 a 20 caracteres
                    print(f"⏭️ Saltando párrafo {i+1}: muy corto ({len(para.text.strip())} caracteres)")
                    continue
                    
                texto = para.text
                L = len(texto)
                
                # Procesar párrafos que tengan al menos 30 caracteres
                if L < 30:
                    print(f"⏭️ Párrafo {i+1} demasiado corto para procesar: {L} caracteres")
                    continue
                
                print(f"🔄 Procesando párrafo {i+1}: {L} caracteres")
                
                # Encontrar posiciones donde insertar los "0" 
                # Si el párrafo es más corto que 'paso', insertar al menos uno en el medio
                if L < paso:
                    # Para párrafos cortos, insertar uno en el medio
                    objetivos = [L // 2] if L > 30 else []
                else:
                    # Para párrafos largos, usar el patrón normal
                    objetivos = list(range(paso, L, paso))
                
                # Lista para guardar las posiciones reales donde insertar
                posiciones_a_reemplazar = []
                
                for obj in objetivos:
                    # Buscar el espacio más cercano
                    izq = texto.rfind(" ", 0, obj + 1)
                    der = texto.find(" ", obj)
                    
                    if izq == -1 and der == -1:
                        continue
                        
                    # Elegir el espacio más cercano
                    if der == -1:
                        idx = izq
                    elif izq == -1:
                        idx = der
                    else:
                        idx = izq if (obj - izq) <= (der - obj) else der
                    
                    # Agregar posición válida
                    if 0 <= idx < len(texto) and texto[idx] == " ":
                        posiciones_a_reemplazar.append(idx)
                
                # Procesar cada posición de atrás hacia adelante
                for pos in reversed(sorted(set(posiciones_a_reemplazar))):
                    try:
                        reemplazar_espacio_por_cero_invisible(para, pos)
                    except Exception as e:
                        print(f"⚠️ Error reemplazando posición {pos}: {e}")
                        # Continuar con las otras posiciones
                        continue
                
                processed_count += 1
                print(f"✅ Párrafo {i+1} procesado ({len(posiciones_a_reemplazar)} reemplazos)")
                
            except Exception as e:
                print(f"⚠️ Error procesando párrafo {i+1}: {e}")
                # Continuar con el siguiente párrafo
                continue

        print(f"📊 Párrafos procesados: {processed_count}/{total_paragraphs}")

        # NUEVO: Hacer todos los "0" invisibles antes de guardar
        hacer_ceros_invisibles(doc)

        # Generar nombre del archivo procesado
        original_path = pathlib.Path(doc_path)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        nuevo_nombre = f"{original_path.stem}_modificado_{timestamp}.docx"
        nuevo_path = os.path.abspath(os.path.join(PROCESSED_FOLDER, nuevo_nombre))
        
        print(f"📁 Guardando en: {nuevo_path}")
        
        # Guardar el documento procesado
        doc.save(nuevo_path)
        
        # Verificar que se guardó correctamente
        if os.path.exists(nuevo_path):
            file_size = os.path.getsize(nuevo_path)
            print(f"✅ Archivo guardado exitosamente: {nuevo_path} ({file_size} bytes)")
        else:
            print(f"❌ Error: El archivo no se guardó: {nuevo_path}")
            raise Exception("El archivo no se pudo guardar correctamente")
        
        return nuevo_path
        
    except Exception as e:
        print(f"❌ Error crítico en procesar_documento:")
        print(f"   📁 Archivo: {doc_path}")
        print(f"   🎯 Paso: {paso}")
        print(f"   💥 Error: {str(e)}")
        print(f"   📍 Tipo: {type(e).__name__}")
        traceback.print_exc()
        raise Exception(f"Error al procesar el documento: {str(e)}")

def reemplazar_espacio_por_cero_invisible(para, posicion):
    """
    Reemplaza un espacio en la posición especificada por un "0" INVISIBLE (color blanco).
    
    Enfoque SIMPLE: insertar el "0" y luego hacer una pasada final para volver invisibles todos los "0".
    """
    texto_original = para.text
    
    if posicion >= len(texto_original) or texto_original[posicion] != " ":
        return
    
    # Encontrar en qué run está el carácter a reemplazar
    char_count = 0
    target_run = None
    pos_in_run = 0
    
    for run in para.runs:
        run_length = len(run.text)
        if char_count <= posicion < char_count + run_length:
            target_run = run
            pos_in_run = posicion - char_count
            break
        char_count += run_length
    
    if target_run is None:
        return
    
    # Verificar que efectivamente hay un espacio en esa posición del run
    if pos_in_run >= len(target_run.text) or target_run.text[pos_in_run] != " ":
        return
    
    # PASO 1: Simplemente cambiar el espacio por "0"
    nuevo_texto = target_run.text[:pos_in_run] + "0" + target_run.text[pos_in_run + 1:]
    target_run.text = nuevo_texto
    
    print(f"✅ Insertado '0' en posición {posicion} (será invisible al final)")

def hacer_ceros_invisibles(doc):
    """
    Función que se ejecuta al final para hacer todos los "0" invisibles en todo el documento.
    """
    print("🎨 Haciendo todos los '0' invisibles...")
    ceros_procesados = 0
    
    for para in doc.paragraphs:
        if "0" in para.text:
            for run in para.runs:
                if "0" in run.text:
                    # Para cada "0" en el run, crear runs separados
                    texto = run.text
                    if "0" in texto:
                        # Reemplazar todo el contenido del run
                        partes = texto.split("0")
                        
                        # Conservar formato original
                        formato_original = {
                            'name': run.font.name,
                            'size': run.font.size,
                            'bold': run.font.bold,
                            'italic': run.font.italic,
                            'underline': run.font.underline
                        }
                        
                        # Conservar color original
                        color_original = None
                        try:
                            if hasattr(run.font.color, 'rgb') and run.font.color.rgb:
                                color_original = run.font.color.rgb
                        except:
                            pass
                        
                        # Limpiar el run actual y poner la primera parte
                        run.text = partes[0]
                        
                        # Insertar las partes restantes con "0" invisibles
                        for i in range(1, len(partes)):
                            # Insertar "0" invisible
                            run_cero = para.add_run("0")
                            run_cero.font.color.rgb = RGBColor(255, 255, 255)  # Blanco = invisible
                            ceros_procesados += 1
                            
                            # Insertar la siguiente parte del texto con formato original
                            if partes[i]:  # Si no está vacía
                                run_normal = para.add_run(partes[i])
                                # Restaurar formato
                                if formato_original['name']:
                                    run_normal.font.name = formato_original['name']
                                if formato_original['size']:
                                    run_normal.font.size = formato_original['size']
                                if formato_original['bold'] is not None:
                                    run_normal.font.bold = formato_original['bold']
                                if formato_original['italic'] is not None:
                                    run_normal.font.italic = formato_original['italic']
                                if formato_original['underline'] is not None:
                                    run_normal.font.underline = formato_original['underline']
                                if color_original:
                                    try:
                                        run_normal.font.color.rgb = color_original
                                    except:
                                        pass
                        break  # Solo procesar el primer run con "0" en cada párrafo
    
    print(f"✅ {ceros_procesados} ceros hechos invisibles (color blanco)")

def copy_run_format(source_font, target_font):
    """
    Copia todo el formato de una fuente a otra.
    """
    try:
        if source_font.name:
            target_font.name = source_font.name
        if source_font.size:
            target_font.size = source_font.size
        if source_font.bold is not None:
            target_font.bold = source_font.bold
        if source_font.italic is not None:
            target_font.italic = source_font.italic
        if source_font.underline is not None:
            target_font.underline = source_font.underline
        # No copiamos el color porque lo haremos invisible
    except:
        # Si hay algún error copiando formato, continuar
        pass

@app.route('/')
def index():
    """Página principal"""
    return render_template('index.html')

@app.route('/health')
def health():
    """Endpoint de salud para verificar que el servidor funciona"""
    return jsonify({"status": "ok", "message": "Servidor funcionando correctamente"})

@app.route('/test')
def test():
    """Endpoint de prueba"""
    return "<h1>¡El servidor Flask está funcionando!</h1><p>La aplicación está corriendo correctamente en Render.</p>"

@app.route('/upload', methods=['POST'])
def upload_file():
    """Manejar la subida y procesamiento de archivos"""
    try:
        print("🔍 Iniciando endpoint /upload...")
        
        # Verificar si se subió un archivo
        if 'file' not in request.files:
            print("❌ No se encontró 'file' en request.files")
            flash('No se seleccionó ningún archivo')
            return redirect(request.url)
        
        file = request.files['file']
        paso = int(request.form.get('paso', 70))
        print(f"📁 Archivo recibido: {file.filename}, paso: {paso}")
        
        # Verificar que se seleccionó un archivo
        if file.filename == '':
            print("❌ Nombre de archivo vacío")
            flash('No se seleccionó ningún archivo')
            return redirect(request.url)
        
        # Verificar que es un archivo permitido
        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"{timestamp}_{filename}"
            filepath = os.path.abspath(os.path.join(UPLOAD_FOLDER, filename))
            print(f"💾 Guardando archivo en: {filepath}")
            
            try:
                file.save(filepath)
                print(f"✅ Archivo guardado: {filepath}")
                
                # Verificar tamaño del archivo guardado
                file_size = os.path.getsize(filepath)
                print(f"📏 Tamaño del archivo: {file_size} bytes")
                
                if file_size == 0:
                    print("❌ El archivo guardado está vacío")
                    os.remove(filepath)
                    return jsonify({
                        'success': False,
                        'message': 'El archivo subido está vacío'
                    }), 400
                
                # Procesar el documento
                print(f"🔄 Iniciando procesamiento de: {filename}")
                processed_path = procesar_documento(filepath, paso)
                print(f"✅ Procesamiento completado: {processed_path}")
                
                # Limpiar archivo temporal
                os.remove(filepath)
                print(f"🗑️ Archivo temporal eliminado: {filepath}")
                
                # Retornar información del archivo procesado
                processed_filename = os.path.basename(processed_path)
                download_url = url_for('download_file', filename=processed_filename)
                
                print(f"📤 URL de descarga generada: {download_url}")
                print(f"📁 Nombre del archivo procesado: {processed_filename}")
                
                return jsonify({
                    'success': True,
                    'message': 'Documento procesado exitosamente',
                    'download_url': download_url,
                    'filename': processed_filename
                })
                
            except Exception as e:
                print(f"❌ Error durante el procesamiento:")
                print(f"   📁 Archivo: {filepath}")
                print(f"   💥 Error: {str(e)}")
                print(f"   🔍 Tipo: {type(e).__name__}")
                import traceback
                traceback.print_exc()
                
                # Limpiar archivo si existe
                try:
                    if os.path.exists(filepath):
                        os.remove(filepath)
                        print(f"🗑️ Archivo temporal limpiado: {filepath}")
                except:
                    pass
                
                # Retornar error específico
                error_msg = f"Error al procesar el documento: {str(e)}"
                if "corrupted" in str(e).lower() or "invalid" in str(e).lower():
                    error_msg = "El archivo parece estar corrupto o tener un formato inválido"
                elif "memory" in str(e).lower() or "size" in str(e).lower():
                    error_msg = "El archivo es demasiado grande para procesar"
                
                return jsonify({
                    'success': False,
                    'message': error_msg
                }), 500
        
        else:
            print(f"❌ Archivo no permitido: {file.filename}")
            flash('Tipo de archivo no permitido. Solo se permiten archivos .docx y .docm')
            return redirect(request.url)
            
    except Exception as e:
        print(f"❌ Error inesperado en /upload: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({
            'success': False,
            'message': f'Error: {str(e)}'
        }), 500

@app.route('/download/<filename>')
def download_file(filename):
    """Descargar archivo procesado"""
    try:
        # Usar ruta absoluta
        file_path = os.path.abspath(os.path.join(PROCESSED_FOLDER, filename))
        print(f"🔍 Intentando descargar: {file_path}")
        print(f"📁 ¿Existe el archivo?: {os.path.exists(file_path)}")
        
        if not os.path.exists(file_path):
            print(f"❌ Archivo no encontrado: {file_path}")
            # Listar archivos en la carpeta para debugging
            try:
                files_in_folder = os.listdir(PROCESSED_FOLDER)
                print(f"📂 Archivos en processed: {files_in_folder}")
            except:
                print("📂 No se pudo listar la carpeta processed")
            return jsonify({'success': False, 'message': 'Archivo no encontrado'}), 404
        
        print(f"✅ Enviando archivo: {filename}")
        return send_file(
            file_path,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        print(f"❌ Error al descargar archivo: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'message': f'Error al descargar: {str(e)}'}), 500

@app.route('/status')
def status():
    """Endpoint para verificar el estado del servidor"""
    return jsonify({'status': 'online', 'message': 'Servidor funcionando correctamente'})

@app.route('/health')
def health():
    """Endpoint de salud para Render.com"""
    return jsonify({
        'status': 'healthy',
        'timestamp': datetime.now().isoformat(),
        'port': os.environ.get('PORT', '5000'),
        'env': os.environ.get('FLASK_ENV', 'development')
    })

@app.route('/test')
def test():
    """Endpoint de prueba simple"""
    return jsonify({
        'message': 'Flask app funcionando correctamente',
        'python_version': sys.version,
        'cwd': os.getcwd(),
        'uploads_dir': UPLOAD_FOLDER,
        'processed_dir': PROCESSED_FOLDER
    })

if __name__ == '__main__':
    print("🚀 Iniciando servidor web...")
    print("📁 Procesador de documentos Word")
    print(f"📦 Puerto configurado: {os.environ.get('PORT', '5000')}")
    print(f"🌍 Entorno: {os.environ.get('FLASK_ENV', 'development')}")
    print(f"📂 Directorio actual: {os.getcwd()}")
    
    # Configuración para deployment
    port = int(os.environ.get('PORT', 5000))
    debug = os.environ.get('FLASK_ENV') != 'production'
    
    # Asegurar que las carpetas existen
    print(f"📁 Creando carpetas: {UPLOAD_FOLDER}, {PROCESSED_FOLDER}")
    os.makedirs(UPLOAD_FOLDER, exist_ok=True)
    os.makedirs(PROCESSED_FOLDER, exist_ok=True)
    
    if debug:
        print("🌐 Accede a: http://localhost:5000")
        print("📝 Sube archivos .docx/.docm para procesarlos")
        print("\n⚡ Presiona Ctrl+C para detener el servidor")
    else:
        print(f"🌐 Servidor iniciando en puerto {port}")
        print("📝 Aplicación en modo producción")
        print("🔍 Endpoints disponibles: /, /status, /health, /test")
    
    try:
        # Configuración mejorada para Render.com
        app.run(
            debug=debug, 
            host='0.0.0.0', 
            port=port,
            threaded=True,
            use_reloader=False
        )
    except Exception as e:
        print(f"❌ Error al iniciar servidor: {e}")
        import traceback
        traceback.print_exc()
